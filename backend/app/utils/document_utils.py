"""
Extract plain text from PDF, DOCX, or TXT files for syllabus parsing.
Also provides structured extraction (sections, tables, candidate dates/percentages) for hybrid LLM parser.
"""
import io
import re
from pathlib import Path


def extract_text_from_file(file) -> str:
    """
    Extract plain text from a file (PDF, DOCX, or TXT).
    file: file-like object (e.g. from request.files) or bytes,
    must have .filename or detect by content.
    Returns plaintext string for parsing.
    """
    # Get filename and content
    filename = getattr(file, "filename", None) or ""
    if hasattr(file, "read"):
        content = file.read()
        file.seek(0)  # Reset for possible reuse
    else:
        content = file

    ext = (Path(filename).suffix or "").lower()
    if ext == ".txt":
        return _extract_txt(content)
    if ext == ".pdf":
        return _extract_pdf(content)
    if ext in (".docx", ".doc"):
        return _extract_docx(content)

    # Fallback: try to detect by magic bytes
    if content[:4] == b"%PDF":
        return _extract_pdf(content)
    # DOCX is a ZIP file
    if content[:2] == b"PK":
        return _extract_docx(content)

    # Default: treat as UTF-8 text
    return _extract_txt(content)


def _extract_txt(content: bytes) -> str:
    """Extract text from plaintext file."""
    return content.decode("utf-8", errors="replace")


def _extract_pdf(content: bytes) -> str:
    """Extract text from PDF using pdfplumber."""
    import pdfplumber

    parts = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
    return "\n\n".join(parts) if parts else ""


def _extract_docx(content: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    from docx import Document

    doc = Document(io.BytesIO(content))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)
    return "\n\n".join(parts) if parts else ""


# --- Structured extraction for hybrid LLM parser ---

SECTION_HEADERS = (
    "grading", "grade", "evaluation", "assessment", "weights",
    "schedule", "calendar", "syllabus", "course information", "course logistics",
    "course overview", "assignments", "homework", "projects", "exams",
    "course outline", "important dates",
)


def extract_structured_from_file(file) -> dict:
    """
    Extract structured content for hybrid parser: sections, tables, candidate dates/percentages.
    Returns dict with keys: raw_text, sections, tables, candidate_dates, candidate_percentages, metadata.
    """
    filename = getattr(file, "filename", None) or ""
    if hasattr(file, "read"):
        content = file.read()
        file.seek(0)
    else:
        content = file

    ext = (Path(filename).suffix or "").lower()
    if ext == ".txt":
        raw_text, tables = _extract_structured_txt(content)
    elif ext == ".pdf":
        raw_text, tables = _extract_structured_pdf(content)
    elif ext in (".docx", ".doc"):
        raw_text, tables = _extract_structured_docx(content)
    else:
        if content[:4] == b"%PDF":
            raw_text, tables = _extract_structured_pdf(content)
        elif content[:2] == b"PK":
            raw_text, tables = _extract_structured_docx(content)
        else:
            raw_text, tables = _extract_structured_txt(content)

    sections = _detect_sections(raw_text)
    candidate_dates = _extract_candidate_dates(raw_text)
    candidate_percentages = _extract_candidate_percentages(raw_text)

    metadata = {
        "source_type": "pdf" if ext == ".pdf" else "docx" if ext in (".docx", ".doc") else "txt",
    }
    return {
        "raw_text": raw_text,
        "sections": sections,
        "tables": tables,
        "candidate_dates": candidate_dates,
        "candidate_percentages": candidate_percentages,
        "metadata": metadata,
    }


def _extract_structured_txt(content: bytes) -> tuple[str, list]:
    raw = content.decode("utf-8", errors="replace")
    return raw, []  # TXT has no table structure


def _extract_structured_pdf(content: bytes) -> tuple[str, list]:
    import pdfplumber
    parts = []
    tables = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
            tbls = page.extract_tables()
            if tbls:
                for t in tbls:
                    if t and any(cell for row in t for cell in (row or []) if cell):
                        tables.append(t)
    raw = "\n\n".join(parts) if parts else ""
    return raw, tables


def _extract_structured_docx(content: bytes) -> tuple[str, list]:
    from docx import Document
    doc = Document(io.BytesIO(content))
    parts = []
    tables = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(cells)
        if rows:
            tables.append(rows)
    raw = "\n\n".join(parts) if parts else ""
    return raw, tables


def _detect_sections(raw_text: str) -> list[dict]:
    """Split text into sections by headers (Grading, Schedule, etc.)."""
    sections = []
    blocks = re.split(r"\n\s*\n+", raw_text)
    current_heading = None
    current_content = []

    for block in blocks:
        lines = block.strip().split("\n")
        if not lines:
            continue
        first = lines[0].strip()
        rest = "\n".join(lines[1:]) if len(lines) > 1 else ""

        # Does first line look like a header? (short, title-like, possibly ends with :)
        is_header = (
            len(first) >= 1
            and len(first) < 60
            and first[0].isalpha()
            and (first.endswith(":") or first.isupper() or first.istitle())
        )
        header_lower = first.lower().rstrip(":")
        if is_header and any(h in header_lower for h in SECTION_HEADERS):
            if current_heading or current_content:
                sections.append({
                    "heading": current_heading or "Content",
                    "content": "\n".join(current_content).strip(),
                })
            current_heading = first.rstrip(":")
            current_content = [rest] if rest else []
        else:
            current_content.append(block)

    if current_heading or current_content:
        sections.append({
            "heading": current_heading or "Content",
            "content": "\n".join(current_content).strip(),
        })

    if not sections:
        sections = [{"heading": "Full document", "content": raw_text[:8000]}]
    return sections


def _extract_candidate_dates(raw_text: str) -> list[dict]:
    """Extract date-like strings with context for LLM."""
    candidates = []
    # Month + day: Oct. 24, November 4, Dec 11 12:30
    pat1 = r"\b(Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\.?\s+\d{1,2}(?:st|nd|rd|th)?(?:\s*,?\s*\d{4})?(?:\s+\d{1,2}:\d{2})?\b"
    # Numeric: 12/11, 12-11-2025
    pat2 = r"\b\d{1,2}[/\-]\d{1,2}(?:[/\-]\d{2,4})?\b"
    # ISO: 2025-10-24, 2025-10-24T12:30
    pat3 = r"\b\d{4}-\d{2}-\d{2}(?:T\d{2}:\d{2})?\b"
    for m in re.finditer(f"({pat1}|{pat2}|{pat3})", raw_text, re.I):
        start = max(0, m.start() - 25)
        end = min(len(raw_text), m.end() + 25)
        context = raw_text[start:end].replace("\n", " ")
        candidates.append({"raw": m.group(0).strip(), "context": context[:80]})
    seen = set()
    out = []
    for c in candidates:
        r = c.get("raw", "")
        if r and r not in seen:
            seen.add(r)
            out.append(c)
    return out[:50]


def _extract_candidate_percentages(raw_text: str) -> list[dict]:
    """Extract percentage mentions with context."""
    out = []
    for m in re.finditer(r".{0,35}(\d{1,3})\s*%.{0,35}", raw_text):
        ctx = m.group(0).replace("\n", " ")
        val = int(re.search(r"(\d{1,3})\s*%", ctx).group(1))
        if 1 <= val <= 100:
            out.append({"value": val, "context": ctx.strip()[:100]})
    seen_ctx = set()
    deduped = []
    for c in out:
        k = (c["value"], c["context"][:50])
        if k not in seen_ctx:
            seen_ctx.add(k)
            deduped.append(c)
    return deduped[:30]
